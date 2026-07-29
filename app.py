import os
import random
from datetime import datetime
import streamlit as st
import io

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image, Flowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.graphics.shapes import Drawing
from reportlab.graphics.barcode.qr import QrCodeWidget
from reportlab.lib.units import cm

from docx import Document

# --- Forceful Seal Overlay Class (Zero-Height Flowable) ---
class OverlaySeal(Flowable):
    def __init__(self, img_path, width=4.5*cm, height=2.5*cm, x_offset=0, y_offset=0, align="center"):
        super().__init__()
        self.img_path = img_path
        self.width = width
        self.height = height
        self.x_offset = x_offset
        self.y_offset = y_offset
        self.align = align

    def wrap(self, availWidth, availHeight):
        # Returns (0, 0) so the document engine treats this as taking 0 height.
        # This guarantees NO unwanted page breaks, even when overlapping text.
        return 0, 0

    def draw(self):
        if os.path.exists(self.img_path):
            self.canv.saveState()
            if self.align == "center":
                x = (570 - self.width) / 2 + self.x_offset
            elif self.align == "right":
                x = 570 - self.width + self.x_offset
            else:
                x = self.x_offset
            
            y = self.y_offset - self.height
            self.canv.drawImage(self.img_path, x, y, width=self.width, height=self.height, mask='auto')
            self.canv.restoreState()


# --- Helper Functions ---
def num_to_words_indian_clean(num):
    num = int(round(num))
    if num == 0: return "ZERO RUPEES ONLY"
    units = ["", "ONE", "TWO", "THREE", "FOUR", "FIVE", "SIX", "SEVEN", "EIGHT", "NINE", "TEN", 
             "ELEVEN", "TWELVE", "THIRTEEN", "FOURTEEN", "FIFTEEN", "SIXTEEN", "SEVENTEEN", "EIGHTEEN", "NINETEEN"]
    tens = ["", "", "TWENTY", "THIRTY", "FORTY", "FIFTY", "SIXTY", "SEVENTY", "EIGHTY", "NINETY"]
    
    def convert_below_thousand(n):
        res = ""
        if n >= 100:
            res += units[n // 100] + " HUNDRED "
            n %= 100
        if n >= 20:
            res += tens[n // 10] + " "
            n %= 10
        if n > 0:
            res += units[n] + " "
        return res.strip()

    result = ""
    crore = num // 10000000; num %= 10000000
    lakh = num // 100000; num %= 100000
    thousand = num // 1000; num %= 1000

    if crore > 0: result += convert_below_thousand(crore) + (" CRORES " if crore > 1 else " CRORE ")
    if lakh > 0: result += convert_below_thousand(lakh) + (" LAKHS " if lakh > 1 else " LAKH ")
    if thousand > 0: result += convert_below_thousand(thousand) + " THOUSAND "
    if num > 0: result += convert_below_thousand(num)
    return f"{result.strip()} RUPEES ONLY"

def generate_pdf_file(pdf_filename, owner, address, est_date, ref_no, target_total, processed_items, item_amounts, actual_subtotal, actual_gst, final_total, is_single_page, include_seal):
    seal_path = "seal_sign.png"
    has_seal = include_seal and os.path.exists(seal_path)

    doc = SimpleDocTemplate(pdf_filename, pagesize=letter, rightMargin=20, leftMargin=20, topMargin=10, bottomMargin=10)
    styles = getSampleStyleSheet()

    RED_COLOR, BLUE_COLOR, LIGHT_PINK, BORDER_BLUE = colors.HexColor("#DC2626"), colors.HexColor("#1E40AF"), colors.HexColor("#EC4899"), colors.HexColor("#2563EB")

    title_style = ParagraphStyle("Title", parent=styles["Heading1"], alignment=1, fontSize=28, leading=32, fontName="Helvetica-Bold", textColor=RED_COLOR)
    sub_style = ParagraphStyle("Sub", parent=styles["Normal"], alignment=1, fontSize=9, leading=12, fontName="Helvetica-Bold", textColor=BLUE_COLOR)
    gstin_style = ParagraphStyle("GSTIN", parent=styles["Normal"], alignment=1, fontSize=9, leading=12, fontName="Helvetica-Bold", textColor=LIGHT_PINK)
    ref_left_style = ParagraphStyle("RefLeft", parent=styles["Normal"], alignment=0, fontSize=10, leading=12, fontName="Helvetica")
    ref_right_style = ParagraphStyle("RefRight", parent=styles["Normal"], alignment=2, fontSize=10, leading=12, fontName="Helvetica")
    box_hdr_style = ParagraphStyle("BoxHdr", parent=styles["Normal"], alignment=1, fontSize=14, leading=16, fontName="Helvetica-Bold", textColor=colors.black)
    box_detail_style = ParagraphStyle("BoxDetail", parent=styles["Normal"], alignment=1, fontSize=12.5, leading=15, fontName="Helvetica-Bold", textColor=colors.black)
    
    if is_single_page:
        cell_12_bold_center = ParagraphStyle("Cell11BC", parent=styles["Normal"], alignment=1, fontSize=11, leading=13.5, fontName="Helvetica-Bold", textColor=colors.black)
        hdr_12_bold_center = ParagraphStyle("Hdr11BC", parent=styles["Normal"], alignment=1, fontSize=11, leading=13.5, fontName="Helvetica-Bold", textColor=colors.black)
        total_14_bold = ParagraphStyle("Total12B", parent=styles["Normal"], alignment=1, fontSize=12, leading=14.5, fontName="Helvetica-Bold", textColor=colors.black)
        words_13_bold_center = ParagraphStyle("Words11.5BC", parent=styles["Normal"], alignment=1, fontSize=11.5, leading=14, fontName="Helvetica-Bold", textColor=colors.black)
        terms_hdr_center = ParagraphStyle("TermsHdr9.5", parent=styles["Normal"], alignment=1, fontSize=9.5, leading=12, fontName="Helvetica-Bold", textColor=colors.black)
        terms_point_size_8 = ParagraphStyle("TermsPt8.5", parent=styles["Normal"], alignment=0, fontSize=8.5, leading=10.5, fontName="Helvetica-Bold", textColor=colors.black)
    else:
        cell_12_bold_center = ParagraphStyle("Cell12BC", parent=styles["Normal"], alignment=1, fontSize=12, leading=14, fontName="Helvetica-Bold", textColor=colors.black)
        hdr_12_bold_center = ParagraphStyle("Hdr12BC", parent=styles["Normal"], alignment=1, fontSize=12, leading=14, fontName="Helvetica-Bold", textColor=colors.black)
        total_14_bold = ParagraphStyle("Total14B", parent=styles["Normal"], alignment=1, fontSize=14, leading=16, fontName="Helvetica-Bold", textColor=colors.black)
        words_13_bold_center = ParagraphStyle("Words13BC", parent=styles["Normal"], alignment=1, fontSize=13, leading=16, fontName="Helvetica-Bold", textColor=colors.black)
        terms_hdr_center = ParagraphStyle("TermsHdr10", parent=styles["Normal"], alignment=1, fontSize=10, leading=14, fontName="Helvetica-Bold", textColor=colors.black)
        terms_point_size_8 = ParagraphStyle("TermsPt8", parent=styles["Normal"], alignment=0, fontSize=8, leading=11, fontName="Helvetica-Bold", textColor=colors.black)

    elements = []

    def create_header_with_qr():
        qr_data = f"CUSTOMER NAME: {owner.upper()}\nADDRESS: {address.upper()}\nREF NO: {ref_no}\nDATE: {est_date}\nESTIMATION AMOUNT: Rs. {final_total:,}"
        qr = QrCodeWidget(qr_data)
        qr_bounds = qr.getBounds()
        w, h = qr_bounds[2] - qr_bounds[0], qr_bounds[3] - qr_bounds[1]
        d = Drawing(60, 60, transform=[60.0/w, 0, 0, 60.0/h, 0, 0])
        d.add(qr)
        header_text_flowables = [
            Paragraph("SND INTERIOR & DESIGNS", title_style), Spacer(1, 2),
            Paragraph("INTERIOR WORKS, DESIGN ESTIMATE, FLOOR VALUATIONS, BUILDING PLANS", sub_style),
            Paragraph("#15, E BLOCK, SAHAKHAR NAGAR, BANGALORE-560092", sub_style),
            Paragraph("GSTIN: 29ABCDE1234F1Z5", gstin_style),
        ]
        header_table = Table([["", header_text_flowables, d]], colWidths=[75, 420, 75])
        header_table.setStyle(TableStyle([('VALIGN', (0,0), (-1,-1), 'MIDDLE'), ('LEFTPADDING', (0,0), (-1,-1), 0), ('RIGHTPADDING', (0,0), (-1,-1), 0)]))
        return [header_table, Spacer(1, 4)]

    elements.extend(create_header_with_qr())
    elements.append(Table([[Paragraph(f"REF NO:-{ref_no}", ref_left_style), Paragraph(f"DATE: {est_date}", ref_right_style)]], colWidths=[285, 285]))
    elements.append(Spacer(1, 4))

    address_parts = [p.strip() for p in address.split(',')]
    mid_idx = len(address_parts) // 2
    addr_line_1 = ", ".join(address_parts[:mid_idx]) if mid_idx > 0 else address
    addr_line_2 = ", ".join(address_parts[mid_idx:]) if mid_idx > 0 else ""

    box_content = [[Paragraph("ESTIMATION FOR RENOVATION & INTERIOR DESIGN WORK AT", box_hdr_style)], [Paragraph("RESIDENTIAL FLAT AT", box_hdr_style)], [Paragraph(addr_line_1.upper(), box_detail_style)]]
    if addr_line_2: box_content.append([Paragraph(addr_line_2.upper(), box_detail_style)])
    box_content.append([Paragraph(f"OWNER: - {owner.upper()}", box_detail_style)])

    project_box = Table(box_content, colWidths=[570])
    project_box.setStyle(TableStyle([('BOX', (0,0), (-1,-1), 2, BORDER_BLUE), ('ROUNDEDCORNERS', [8, 8, 8, 8]), ('TOPPADDING', (0,0), (-1,-1), 3), ('BOTTOMPADDING', (0,0), (-1,-1), 3)]))
    elements.append(project_box)
    elements.append(Spacer(1, 5))

    if is_single_page:
        p_table_data = [[Paragraph("SL.NO", hdr_12_bold_center), Paragraph("Description", hdr_12_bold_center), Paragraph("Qty", hdr_12_bold_center), Paragraph("Amount Rs.", hdr_12_bold_center)]]
        for idx in range(10):
            item = processed_items[idx]
            p_table_data.append([Paragraph(f"{idx+1}.", cell_12_bold_center), Paragraph(item[0], cell_12_bold_center), Paragraph(item[1], cell_12_bold_center), Paragraph(f"{item_amounts[idx]:,}", cell_12_bold_center)])
        p_table_data.append(["", Paragraph("GST 18%", total_14_bold), "", Paragraph(f"{actual_gst:,}", total_14_bold)])
        p_table_data.append(["", Paragraph("TOTAL", total_14_bold), "", Paragraph(f"{final_total:,}", total_14_bold)])

        t1 = Table(p_table_data, colWidths=[50, 300, 100, 120])
        t1.setStyle(TableStyle([('GRID', (0,0), (-1,-1), 0.5, colors.black), ('VALIGN', (0,0), (-1,-1), 'MIDDLE'), ('TOPPADDING', (0,0), (-1,-1), 5), ('BOTTOMPADDING', (0,0), (-1,-1), 5)]))
        elements.append(t1)

        elements.append(Spacer(1, 5))
        elements.append(Paragraph(num_to_words_indian_clean(final_total), words_13_bold_center))
        elements.append(Spacer(1, 4))
        elements.append(Paragraph("TERMS AND CONDITIONS:", terms_hdr_center))
        elements.append(Spacer(1, 2))

        terms_points = [
            "1. This Is A Preliminary Estimate And Not A Final Invoice",
            "2. Payment: 50% Advance, 30% After Material Delivery, 20% Upon Completion.",
            "3. Validity: This Estimation Is Valid For 45 Days From The Date Of Issue.",
            "4. Scope Of Work: Any Work Not Explicitly Mentioned In This Estimate Will Be Charged Extra.",
            "5. Materials: All Materials Used Will Be Of Standard Quality Unless Specified Otherwise.",
            "6. Project Duration: Estimated Project Completion Time Is 90 Working Days From Advance."
        ]
        for p_text in terms_points:
            elements.append(Paragraph(p_text, terms_point_size_8))
            elements.append(Spacer(1, 1.5))

        # Single Page Seal: Forcefully pasted at bottom right corner
        if has_seal:
            elements.append(OverlaySeal(seal_path, width=4.5*cm, height=2.5*cm, align="right", y_offset=15))

    else:
        # --- Page 1 of Two-Page Estimation ---
        p1_table_data = [[Paragraph("SL.NO", hdr_12_bold_center), Paragraph("Description", hdr_12_bold_center), Paragraph("Qty", hdr_12_bold_center), Paragraph("Amount Rs.", hdr_12_bold_center)]]
        for idx in range(9):
            item = processed_items[idx]
            p1_table_data.append([Paragraph(f"{idx+1}.", cell_12_bold_center), Paragraph(item[0], cell_12_bold_center), Paragraph(item[1], cell_12_bold_center), Paragraph(f"{item_amounts[idx]:,}", cell_12_bold_center)])
        
        t1 = Table(p1_table_data, colWidths=[50, 300, 100, 120])
        # Standard padding kept at 10 (no reduction)
        t1.setStyle(TableStyle([('GRID', (0,0), (-1,-1), 0.5, colors.black), ('VALIGN', (0,0), (-1,-1), 'MIDDLE'), ('TOPPADDING', (0,0), (-1,-1), 10), ('BOTTOMPADDING', (0,0), (-1,-1), 10)]))
        elements.append(t1)

        # Page 1 Seal: Forcefully pasted at bottom center of Page 1
        if has_seal:
            elements.append(OverlaySeal(seal_path, width=4.5*cm, height=2.5*cm, align="center", y_offset=10))

        elements.append(PageBreak())
        elements.extend(create_header_with_qr())

        # --- Page 2 of Two-Page Estimation ---
        p2_table_data = [[Paragraph("SL.NO", hdr_12_bold_center), Paragraph("Description", hdr_12_bold_center), Paragraph("Qty", hdr_12_bold_center), Paragraph("Amount Rs.", hdr_12_bold_center)]]
        for idx in range(9, 15):
            item = processed_items[idx]
            p2_table_data.append([Paragraph(f"{idx+1}.", cell_12_bold_center), Paragraph(item[0], cell_12_bold_center), Paragraph(item[1], cell_12_bold_center), Paragraph(f"{item_amounts[idx]:,}", cell_12_bold_center)])

        p2_table_data.append(["", Paragraph("GST 18%", total_14_bold), "", Paragraph(f"{actual_gst:,}", total_14_bold)])
        p2_table_data.append(["", Paragraph("TOTAL", total_14_bold), "", Paragraph(f"{final_total:,}", total_14_bold)])

        t2 = Table(p2_table_data, colWidths=[50, 300, 100, 120])
        # Standard padding kept at 10
        t2.setStyle(TableStyle([('GRID', (0,0), (-1,-1), 0.5, colors.black), ('VALIGN', (0,0), (-1,-1), 'MIDDLE'), ('TOPPADDING', (0,0), (-1,-1), 10), ('BOTTOMPADDING', (0,0), (-1,-1), 10)]))
        elements.append(t2)

        elements.append(Spacer(1, 6))
        elements.append(Paragraph(num_to_words_indian_clean(final_total), words_13_bold_center))
        elements.append(Spacer(1, 5))
        elements.append(Paragraph("TERMS AND CONDITIONS:", terms_hdr_center))
        elements.append(Spacer(1, 3))

        terms_points = [
            "1. This Is A Preliminary Estimate And Not A Final Invoice",
            "2. Payment: 50% Advance, 30% After Material Delivery, 20% Upon Completion.",
            "3. Validity: This Estimation Is Valid For 45 Days From The Date Of Issue.",
            "4. Scope Of Work: Any Work Not Explicitly Mentioned In This Estimate Will Be Charged Extra.",
            "5. Materials: All Materials Used Will Be Of Standard Quality Unless Specified Otherwise.",
            "6. Project Duration: Estimated Project Completion Time Is 90 Working Days From Advance."
        ]
        for point in terms_points:
            elements.append(Paragraph(point, terms_point_size_8))
            elements.append(Spacer(1, 2))

        # Page 2 Seal: Forcefully pasted at bottom center of Page 2
        if has_seal:
            elements.append(OverlaySeal(seal_path, width=4.5*cm, height=2.5*cm, align="center", y_offset=-5))

    doc.build(elements)


def generate_estimation_pdf_both(owner, address, est_date, target_total):
    output_dir = "ESTIMATIONS"
    os.makedirs(output_dir, exist_ok=True)

    is_single_page = target_total < 1500000

    FIXED_ITEMS_MASTER = [
        ("Replacing sanitary fittings inside the toilets", "SETS", "SETS_UNITS", 0.08),
        ("3 course of oil bond distemper (Inside repaint)", "SQ. FT", "SQFT", 0.07),
        ("Providing & casting bathroom glazed tiles fixing etc.", "SQ. FT", "SQFT", 0.05),
        ("Interior works (Wardrobes, Modular Kitchen)", "JOB", "JOB_LOT", 0.12),
        ("Electrical fittings, cables, switches etc.", "JOB", "JOB_LOT", 0.07),
        ("Painting (exterior walls)", "SQ. FT", "SQFT", 0.06),
        ("New plumbing lines and fixtures", "JOB", "JOB_LOT", 0.05),
        ("Landscaping/Balcony improvements", "JOB", "JOB_LOT", 0.06),
        ("False ceiling work", "SQ. FT", "SQFT", 0.07),
        ("Flooring (Tiles/Marble)", "SQ. FT", "SQFT", 0.07),
        ("Providing & fixing teak wood show case", "UNIT", "SETS_UNITS", 0.06),
        ("2 course of snow cem paint (Outside repaint)", "SQ. FT", "SQFT", 0.04),
        ("Replacing sanitary fittings inside the kitchen", "SET", "SETS_UNITS", 0.05),
        ("Demolition and debris removal", "LOT", "JOB_LOT", 0.06),
        ("Wall plastering and finishing", "SQ. FT", "SQFT", 0.09)
    ]

    def calculate_quantity(category, total_amount):
        min_budget, max_budget = 1500000.0, 4500000.0
        ratio = max(0.0, min(1.0, (total_amount - min_budget) / (max_budget - min_budget)))
        ratio = max(0.0, min(1.0, ratio + random.uniform(-0.05, 0.05)))
        if category == "SQFT": return f"{round(650 + ratio * (2500 - 650))} SQ. FT"
        elif category == "SETS_UNITS":
            qty = round(1 + ratio * (5 - 1))
            return f"{qty} SETS" if qty > 1 else "1 SET"
        elif category == "JOB_LOT":
            qty = round(1 + ratio * (2 - 1))
            return f"{qty} JOB" if qty > 1 else "1 JOB"
        return "1 JOB"

    total_items_needed = 10 if is_single_page else 15
    processed_items = [(desc, calculate_quantity(cat, target_total), w) for desc, _, cat, w in FIXED_ITEMS_MASTER]
    random.shuffle(processed_items)
    processed_items = processed_items[:total_items_needed]

    subtotal_target = target_total / 1.18
    weights = [item[2] * random.uniform(0.85, 1.15) for item in processed_items]
    total_weight = sum(weights)
    norm_weights = [w / total_weight for w in weights]
    item_amounts = [round(subtotal_target * w) for w in norm_weights]
    item_amounts[-1] += round(subtotal_target) - sum(item_amounts)

    actual_subtotal = sum(item_amounts)
    actual_gst = round(actual_subtotal * 0.18)
    final_total = actual_subtotal + actual_gst

    now = datetime.now()
    ref_no = now.strftime("%H%M%d%m%Y")
    clean_owner_name = owner.replace(' ', '_').replace('&', 'AND')

    sealed_pdf_path = os.path.join(output_dir, f"Estimation_{ref_no}_{clean_owner_name}_Sealed.pdf")
    unsealed_pdf_path = os.path.join(output_dir, f"Estimation_{ref_no}_{clean_owner_name}_Unsealed.pdf")

    # Generate Sealed PDF
    generate_pdf_file(sealed_pdf_path, owner, address, est_date, ref_no, target_total, processed_items, item_amounts, actual_subtotal, actual_gst, final_total, is_single_page, include_seal=True)
    
    # Generate Unsealed PDF
    generate_pdf_file(unsealed_pdf_path, owner, address, est_date, ref_no, target_total, processed_items, item_amounts, actual_subtotal, actual_gst, final_total, is_single_page, include_seal=False)

    # Logging
    docx_filename = os.path.join(output_dir, "ESTIMATION_LOG.docx")
    if os.path.exists(docx_filename):
        doc_word = Document(docx_filename)
    else:
        doc_word = Document()
        doc_word.add_heading('ESTIMATION LOGS', level=1)
        table = doc_word.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text, hdr_cells[3].text = 'REF NO', 'DATE', 'CUSTOMER NAME', 'AMOUNT (Rs.)'

    table = doc_word.tables[0]
    row_cells = table.add_row().cells
    row_cells[0].text, row_cells[1].text, row_cells[2].text, row_cells[3].text = str(ref_no), str(est_date), str(owner.upper()), f"{final_total:,.2f}"
    doc_word.save(docx_filename)

    return sealed_pdf_path, unsealed_pdf_path, docx_filename, ref_no


# --- Streamlit Web UI ---
st.set_page_config(page_title="SND Estimation Generator", page_icon="🏗️", layout="wide")

st.title("🏗️ SND Interior & Designs")
st.subheader("Estimation PDF Generator")

tabs = st.tabs(["✨ Generate New Estimation", "🔍 Re-download Estimation by Ref No"])

# --- TAB 1: GENERATE NEW ESTIMATION ---
with tabs[0]:
    with st.form("estimation_form"):
        st.markdown("### Enter Customer & Project Details")
        owner_input = st.text_input("Owner Name:", value="KUSHAL ANAND & HKS")
        address_input = st.text_area("Site Address:", value="BIRLA TRIMAYA PHASE 4 FLAT-1205, T-6, F-12, DEVANAHALLI CHIKKAJALA, BENGALURU")
        date_input = st.text_input("Date:", value=datetime.now().strftime("%d-%m-%Y"))
        amount_input = st.number_input("Target Amount (Rs.):", min_value=50000, max_value=10000000, value=1499000, step=10000)
        
        include_seal_choice = st.checkbox("Include Seal & Signature by default", value=True)
        
        submitted = st.form_submit_button("Generate PDF", type="primary")

    if submitted:
        with st.spinner('Calculating items and generating Sealed & Unsealed PDFs...'):
            try:
                sealed_path, unsealed_path, docx_path, generated_ref = generate_estimation_pdf_both(
                    owner_input, address_input, date_input, float(amount_input)
                )
                st.success(f"✅ Estimation Generated Successfully! (Reference No: `{generated_ref}`)")
                
                col1, col2, col3 = st.columns(3)
                
                # Download Sealed PDF
                with open(sealed_path, "rb") as file_sealed:
                    col1.download_button(
                        label="📄 Download Sealed PDF",
                        data=file_sealed,
                        file_name=os.path.basename(sealed_path),
                        mime="application/pdf"
                    )

                # Download Unsealed PDF
                with open(unsealed_path, "rb") as file_unsealed:
                    col2.download_button(
                        label="📄 Download Unsealed PDF",
                        data=file_unsealed,
                        file_name=os.path.basename(unsealed_path),
                        mime="application/pdf"
                    )

                # Download Log
                with open(docx_path, "rb") as file_log:
                    col3.download_button(
                        label="📋 Download Log (Word)",
                        data=file_log,
                        file_name="ESTIMATION_LOG.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                st.info("💡 **Tip:** Save the Reference Number (`" + generated_ref + "`) to re-download these PDFs anytime from the Re-download tab.")

            except Exception as e:
                st.error(f"An error occurred: {e}")

# --- TAB 2: RE-DOWNLOAD ESTIMATION BY REFERENCE NUMBER ---
with tabs[1]:
    st.markdown("### Search & Re-download Existing Estimations")
    search_ref = st.text_input("Enter Reference Number (e.g., 182329072026):").strip()
    
    if st.button("🔍 Find Estimation", type="primary"):
        if not search_ref:
            st.warning("Please enter a valid Reference Number.")
        else:
            output_dir = "ESTIMATIONS"
            found_sealed = None
            found_unsealed = None

            if os.path.exists(output_dir):
                for filename in os.listdir(output_dir):
                    if search_ref in filename:
                        full_path = os.path.join(output_dir, filename)
                        if "Sealed" in filename:
                            found_sealed = full_path
                        elif "Unsealed" in filename:
                            found_unsealed = full_path

            if found_sealed or found_unsealed:
                st.success(f"✅ Found Estimation files for Reference No: `{search_ref}`")
                r_col1, r_col2 = st.columns(2)
                
                if found_sealed and os.path.exists(found_sealed):
                    with open(found_sealed, "rb") as f_s:
                        r_col1.download_button(
                            label="📄 Re-Download Sealed PDF",
                            data=f_s,
                            file_name=os.path.basename(found_sealed),
                            mime="application/pdf"
                        )
                else:
                    r_col1.warning("Sealed PDF file not found.")

                if found_unsealed and os.path.exists(found_unsealed):
                    with open(found_unsealed, "rb") as f_u:
                        r_col2.download_button(
                            label="📄 Re-Download Unsealed PDF",
                            data=f_u,
                            file_name=os.path.basename(found_unsealed),
                            mime="application/pdf"
                        )
                else:
                    r_col2.warning("Unsealed PDF file not found.")
            else:
                st.error(f"❌ No estimation found with Reference Number: `{search_ref}`. Please check the number and try again.")
